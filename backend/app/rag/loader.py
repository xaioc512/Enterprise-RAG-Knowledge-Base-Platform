"""文档加载器 — 支持 PDF / Word / Markdown / TXT"""

from pathlib import Path
from loguru import logger


class DocumentLoadError(Exception):
    """文档加载异常"""
    pass


def load_pdf(file_path: str) -> str:
    """加载 PDF 文档，逐页提取文本"""
    try:
        from PyPDF2 import PdfReader
        reader = PdfReader(file_path)
        pages = []
        for i, page in enumerate(reader.pages):
            text = page.extract_text()
            if text:
                pages.append(text)
        content = "\n\n".join(pages)
        if not content.strip():
            raise DocumentLoadError("PDF 中未提取到文本内容（可能是扫描件）")
        logger.info(f"PDF loaded: {len(pages)} pages, {len(content)} chars")
        return content
    except ImportError:
        raise DocumentLoadError("PyPDF2 未安装")


def load_docx(file_path: str) -> str:
    """加载 Word 文档，逐段提取文本"""
    try:
        from docx import Document
        doc = Document(file_path)
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
        # 也提取表格中的文本
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        paragraphs.append(cell.text.strip())
        content = "\n\n".join(paragraphs)
        logger.info(f"DOCX loaded: {len(paragraphs)} paragraphs, {len(content)} chars")
        return content
    except ImportError:
        raise DocumentLoadError("python-docx 未安装")


def load_markdown(file_path: str) -> str:
    """加载 Markdown 文档"""
    try:
        import markdown
        with open(file_path, "r", encoding="utf-8") as f:
            raw = f.read()
        # 转为纯文本（去除 HTML 标签）
        html = markdown.markdown(raw)
        # 简单去除 HTML 标签
        import re
        text = re.sub(r"<[^>]+>", "", html)
        logger.info(f"Markdown loaded: {len(raw)} chars raw, {len(text)} chars text")
        return text.strip()
    except ImportError:
        raise DocumentLoadError("markdown 未安装")


def load_txt(file_path: str) -> str:
    """加载纯文本文档"""
    with open(file_path, "r", encoding="utf-8", errors="replace") as f:
        content = f.read()
    logger.info(f"TXT loaded: {len(content)} chars")
    return content.strip()


LOADERS = {
    "pdf": load_pdf,
    "docx": load_docx,
    "md": load_markdown,
    "txt": load_txt,
}


def load_document(file_path: str, file_type: str) -> str:
    """根据文件类型加载文档，返回纯文本内容"""
    loader = LOADERS.get(file_type)
    if not loader:
        raise DocumentLoadError(f"不支持的文件类型: {file_type}")

    path = Path(file_path)
    if not path.exists():
        raise DocumentLoadError(f"文件不存在: {file_path}")

    return loader(str(path))
